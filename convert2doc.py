import os
from docx import Document

#i'm enjoying these emojis with my print statements 
def combine_folder_to_docx(folder_path, output_folder, doc_name):
    """
    Combines all .txt files in a folder into a single .docx file.
    """
    try:
        # Create output folder if it doesn't exist
        os.makedirs(output_folder, exist_ok=True)

        # Create the .docx document
        doc = Document()

        # Iterate over all .txt files and append to the .docx
        txt_files = [f for f in os.listdir(folder_path) if f.endswith('.txt')]
        
        # Sort to maintain file order
        txt_files.sort()

        if not txt_files:
            print(f"⚠️ No .txt files found in {folder_path}. Skipping...")
            return

        print(f"📄 Combining {len(txt_files)} files from {folder_path} → {doc_name}.docx")

        for text_file in txt_files:
            file_path = os.path.join(folder_path, text_file)

            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()

            # Skip empty files
            if not content.strip():
                print(f"⚠️ Skipping empty file: {file_path}")
                continue

            # Add content to the .docx
            doc.add_paragraph(f"--- {text_file} ---\n")
            doc.add_paragraph(content)
            doc.add_page_break()  # Add page break between files

        # Save the combined .docx
        output_file = os.path.join(output_folder, f"{doc_name}.docx")
        doc.save(output_file)
        print(f"✅ Created: {output_file}")

    except Exception as e:
        print(f"❌ Error: {e}")


# 🔥 **Main Execution**
input_folder = "/Users/shhreya/hindi-ocr/outputtext"      # Path to input .txt folders
output_folder = "/Users/shhreya/hindi-ocr/combined_docs"  # Folder for final .docx files

# List of folders to combine
folders = ["introduction", "part1:4", "part2:4", "part3:4", "part4:4"]

# Combine each folder into a single docx
for folder in folders:
    folder_path = os.path.join(input_folder, folder)
    
    if os.path.exists(folder_path):
        combine_folder_to_docx(folder_path, output_folder, folder)
    else:
        print(f"⚠️ Folder not found: {folder_path}")

print("\n✅ All folders combined successfully!")
